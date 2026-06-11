from __future__ import annotations

import math
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT_DIR = Path.home() / "Documents" / "CAD_Output"
OUT_DIR.mkdir(parents=True, exist_ok=True)


@dataclass(frozen=True)
class Spec:
    material: str = "10钢"
    thickness: float = 2.0
    a1: float = 42.0
    a2: float = 46.0
    b: float = 25.0
    c: float = 30.0
    d: float = 57.0
    e: float = 13.0
    r: float = 5.0
    batch: str = "小批量"
    tolerance: str = "未注公差按 IT14"

    @property
    def neutral_radius(self) -> float:
        return self.r + 0.5 * self.thickness

    @property
    def bend_allowance_90(self) -> float:
        return math.pi / 2 * self.neutral_radius

    @property
    def flat_length(self) -> float:
        bottom_straight = max(self.a1 - 2 * self.r, 0)
        side_straight = max(self.b - self.r, 0)
        return bottom_straight + 2 * side_straight + 2 * self.bend_allowance_90

    @property
    def blank_area(self) -> float:
        return self.flat_length * self.d

    @property
    def hole_area_total(self) -> float:
        return 2 * math.pi * (self.e / 2) ** 2

    @property
    def net_blank_area(self) -> float:
        return self.blank_area - self.hole_area_total


def md_report(spec: Spec) -> str:
    bend_force_note = (
        "弯曲力按课设常用公式估算，可采用 F = k * b * t^2 * sigma_b / (r + t) "
        "或按教材/手册给定 U 形件弯曲公式复核。由于本验证未锁定材料强度表，"
        "此处不强行给出最终吨位，只给出计算路径。"
    )
    return f"""# U形支架冲压工艺及弯曲模具设计说明书初稿

生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}

## 1 设计任务

本课题为 U 形支架冲压工艺及其弯曲模具设计。零件材料为 {spec.material}，材料厚度 t={spec.thickness:.1f} mm，生产批量为{spec.batch}，{spec.tolerance}。

根据题目文档和题图，可确认零件为带两孔的 U 形弯曲件，两处弯曲圆角标注为 2-R5。题目文字中明确给出 A1={spec.a1:.0f} mm，B={spec.b:.0f} mm，C={spec.c:.0f} mm，D={spec.d:.0f} mm。A2 和 E 在原始文档/题图中存在不完整或需人工复核之处，本验证中暂取 A2={spec.a2:.0f} mm，E={spec.e:.0f} mm nominal。

## 2 零件工艺性分析

该零件为低碳钢板料弯曲件，材料塑性较好，适合冷冲压成形。零件主要特征包括：矩形底板、两个圆孔、两侧 U 形弯边和 R5 圆角。由于孔位位于底板中心区域，通常应先完成冲孔和外形落料，再进行 U 形弯曲，以避免弯曲后冲孔定位困难。

从批量看，题目为小批量生产，模具结构宜采用简单、可靠、易加工的单工序或组合工序方案，不宜采用复杂级进模。课程设计中可采用“冲孔落料复合或单工序制坯 + U 形弯曲模”的路线。

## 3 工艺方案比较

方案一：先落料，再冲孔，再弯曲。优点是各工序简单，模具制造容易；缺点是工序数量多，定位累积误差较大。

方案二：冲孔与落料合并制坯，再 U 形弯曲。优点是孔与外形相对位置精度较高，生产效率较好；缺点是制坯模具结构比单工序略复杂。

方案三：采用级进模连续完成冲孔、切边、弯曲。优点是效率高；缺点是模具复杂、成本高，不适合小批量课程设计。

综合考虑生产批量和课程设计要求，推荐采用方案二：冲孔/落料制坯，然后 U 形弯曲。

## 4 毛坯展开尺寸估算

弯曲半径 R={spec.r:.1f} mm，材料厚度 t={spec.thickness:.1f} mm。按课设常用近似，取中性层半径：

Rn = R + 0.5t = {spec.r:.1f} + 0.5×{spec.thickness:.1f} = {spec.neutral_radius:.1f} mm

单个 90° 弯曲展开长度：

L弯 = π/2 × Rn = {spec.bend_allowance_90:.2f} mm

底部直线段近似取 A1 - 2R = {spec.a1:.0f} - 2×{spec.r:.0f} = {spec.a1 - 2 * spec.r:.2f} mm

两侧直线段近似取 2×(B - R) = 2×({spec.b:.0f} - {spec.r:.0f}) = {2 * (spec.b - spec.r):.2f} mm

估算毛坯展开长度：

L = (A1 - 2R) + 2(B - R) + 2L弯 = {spec.flat_length:.2f} mm

毛坯宽度按 D={spec.d:.0f} mm，估算毛坯面积：

A = L × D = {spec.flat_length:.2f} × {spec.d:.0f} = {spec.blank_area:.2f} mm²

两孔面积：

A孔 = 2 × π × (E/2)^2 = {spec.hole_area_total:.2f} mm²

扣孔后净面积约：

A净 = {spec.net_blank_area:.2f} mm²

## 5 冲压工序力计算路径

冲孔力按下式估算：

F冲 = L周 × t × τ

其中 L周 = 2πE，为两个圆孔的总剪切周长，t 为板厚，τ 为材料抗剪强度。10钢的强度参数应按课程教材或冲压手册取值。

落料力按下式估算：

F落 = L外 × t × τ

其中 L外 为毛坯外轮廓周长。若采用矩形毛坯，可按 2(L + D) 估算。

弯曲力按教材 U 形件弯曲公式或常用简化公式估算。{bend_force_note}

压力机公称压力应大于各工序最大压力，并考虑卸料力、顶件力和安全系数。课程设计可初选开式可倾压力机或小吨位机械压力机，再校核闭合高度、工作台尺寸和滑块行程。

## 6 模具结构方案

弯曲模采用简单 U 形弯曲模结构。上模部分包括上模座、垫板、凸模固定板和 U 形弯曲凸模；下模部分包括下模座、凹模块、定位元件和导向元件。工件放置在凹模上，由凸模下行压入凹模型腔完成两侧弯曲。

为保证孔位和弯曲边相对位置，应在制坯阶段保证孔与外形定位精度，弯曲模中可利用外形边或孔进行辅助定位。由于为小批量生产，定位结构以简单可靠为主。

## 7 主要零件设计说明

### 7.1 U形弯曲凸模

凸模工作部分宽度按 A1={spec.a1:.0f} mm 初取，工作圆角取 R5。凸模材料可选 T10A、Cr12 或课程设计常用模具钢，工作部分需热处理。

### 7.2 凹模块

凹模型腔宽度按制件外形和回弹补偿确定。本验证图中按 A2 加间隙余量示意，实际设计应结合回弹、材料厚度和课程手册修正。

### 7.3 定位与导向

总装草图中示意设置导柱导套，以保证上下模相对运动精度。课程设计可按标准模架选用导柱导套规格。

## 8 当前自动生成图纸

当前已生成三张 AutoCAD DWG：

1. `u_bracket_bending_die_validation_20260611_153033.dwg`：零件视图和初始弯曲模剖面。
2. `u_bracket_bending_die_assembly_20260611_154407.dwg`：弯曲模总装草图。
3. `u_bracket_die_detail_set_20260611_155925.dwg`：毛坯展开、凸模、凹模块和计算快照。

## 9 待复核项

1. A2 尺寸需以原始清晰题图或教师给定尺寸为准。
2. E 尺寸在题目文本中显示为 E=13mm+，需确认是否带上偏差或其他标注。
3. 材料强度参数需按课程教材表格取值。
4. 压力机型号需结合最终冲裁力、弯曲力和设备表确定。
5. 图纸若用于正式提交，还需按学校制图规范补图框、标题栏、比例和明细栏。
"""


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.add_run(text)
    return p


def build_docx(markdown_text: str, out_docx: Path):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("U形支架冲压工艺及弯曲模具设计说明书初稿")
    run.bold = True
    run.font.size = Pt(16)

    for raw in markdown_text.splitlines()[2:]:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:], 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 2)
        elif line.startswith("# "):
            continue
        elif line.startswith("- ") or line[0:2].isdigit():
            doc.add_paragraph(line, style=None)
        else:
            add_para(doc, line)

    doc.save(out_docx)


def main():
    spec = Spec()
    text = md_report(spec)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    md_path = OUT_DIR / f"u_bracket_design_report_draft_{stamp}.md"
    docx_path = OUT_DIR / f"u_bracket_design_report_draft_{stamp}.docx"
    md_path.write_text(text, encoding="utf-8")
    build_docx(text, docx_path)
    print(f"md={md_path}")
    print(f"docx={docx_path}")


if __name__ == "__main__":
    main()
